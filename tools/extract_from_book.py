"""extract_from_book.py

Утилита для полуавтоматического извлечения фрагментов из docx-учебника.
Цель: быстро находить разделы (главы/подзаголовки) и собирать черновики
для `webapp/src/content/*.json`.

Использование:
  python tools/extract_from_book.py --docx "Детская психиатрия 2.docx" --list-chapters
  python tools/extract_from_book.py --docx "Детская психиатрия 2.docx" --chapter 7 --out out_ch7.txt
  python tools/extract_from_book.py --docx "Детская психиатрия 2.docx" --search "атомоксетин" --out hits.txt

Скрипт ничего не «решает» автоматически: он помогает быстро найти нужные места,
чтобы вы руками (или вместе со мной) упаковали их в короткие пункты для родителей.
"""

import argparse
import re
from pathlib import Path
from docx import Document

def load_doc(path: str) -> Document:
    p = Path(path)
    if not p.exists():
        raise SystemExit(f"Docx not found: {p}")
    return Document(str(p))

def iter_paragraphs(doc: Document):
    for i, p in enumerate(doc.paragraphs):
        yield i, (p.text or "").strip()

def find_real_chapter_starts(doc: Document):
    starts = []
    for i, t in iter_paragraphs(doc):
        if not t.startswith("Глава "):
            continue
        if "\t" in t:
            continue
        if not re.match(r"^Глава\s+(\d+)\.", t):
            continue
        # эвристика: следующая строка "Введение"
        if i + 1 < len(doc.paragraphs) and (doc.paragraphs[i+1].text or "").strip() == "Введение":
            m = re.match(r"^Глава\s+(\d+)\.(.*)$", t)
            if m:
                num = int(m.group(1))
                title = t
                starts.append((num, title, i))
    starts.sort(key=lambda x: x[0])
    return starts

def chapter_range(starts, num, total_paras):
    # starts: list of (num,title,idx)
    idx = None
    for n, _, i in starts:
        if n == num:
            idx = i
            break
    if idx is None:
        raise SystemExit(f"Chapter {num} not found.")
    # next start
    next_idx = total_paras
    for n, _, i in starts:
        if i > idx:
            next_idx = i
            break
    return idx, next_idx

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True, help="Path to .docx")
    ap.add_argument("--list-chapters", action="store_true", help="List chapter numbers and titles")
    ap.add_argument("--chapter", type=int, help="Extract a whole chapter by number")
    ap.add_argument("--search", type=str, help="Search for keyword/regex and output hits with context")
    ap.add_argument("--out", type=str, help="Output file (txt)")
    ap.add_argument("--context", type=int, default=2, help="Context lines for search")
    args = ap.parse_args()

    doc = load_doc(args.docx)
    starts = find_real_chapter_starts(doc)

    if args.list_chapters:
        for n, title, i in starts:
            print(f"{n:02d}  @para {i:5d}  {title}")
        return

    if args.chapter is not None:
        start, end = chapter_range(starts, args.chapter, len(doc.paragraphs))
        lines = []
        for i in range(start, end):
            t = (doc.paragraphs[i].text or "").rstrip()
            if t:
                lines.append(t)
        out = "\n".join(lines)
        if args.out:
            Path(args.out).write_text(out, encoding="utf-8")
        else:
            print(out)
        return

    if args.search:
        pat = re.compile(args.search, flags=re.IGNORECASE)
        hits = []
        for i, t in iter_paragraphs(doc):
            if not t:
                continue
            if pat.search(t):
                start = max(0, i - args.context)
                end = min(len(doc.paragraphs), i + args.context + 1)
                snippet = []
                for j in range(start, end):
                    tt = (doc.paragraphs[j].text or "").rstrip()
                    if tt:
                        snippet.append(f"{j:5d}: {tt}")
                hits.append("\n".join(snippet))
        out = "\n\n---\n\n".join(hits[:200])
        if args.out:
            Path(args.out).write_text(out, encoding="utf-8")
        else:
            print(out)
        return

    ap.print_help()

if __name__ == "__main__":
    main()
